import uvicorn
import os
import httpx
import json
from io import BytesIO
from typing import List, Optional
from urllib.parse import urlparse
from collections import defaultdict
from itertools import groupby

import trafilatura
from openai import AsyncOpenAI
from pydantic import BaseModel
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.concurrency import run_in_threadpool
from fastapi.responses import StreamingResponse

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.shared import Pt, RGBColor

# --- Configuration ---
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)

MAX_CONTENT_LENGTH = 5 * 1024 * 1024  # 5 MB
TIMEOUT = 15.0  

# --- Data Models ---
class AnalyzeRequest(BaseModel):
    url: str
    api_key: str
    model: str = "gpt-5.4"
    raw_text: Optional[str] = None
    available_categories: Optional[List[str]] = None
    language: str = "tr"  # YENİ EKLENEN DİL PARAMETRESİ


class AnalyzeResponse(BaseModel):
    headline: str
    source: str
    summary: str
    expat_significance: str
    score: int
    category: str
    final_url: Optional[str] = None


class ExportItem(BaseModel):
    input_url: str
    analysis: AnalyzeResponse


# --- Word Helpers ---
def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0563C1')
    rPr.append(color)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink


def ensure_sentence_ends(text: str) -> str:
    t = (text or "").strip()
    if not t: return ""
    if t[-1] in ".!?…": return t
    return t + "."

# --- Security & Fetching ---
async def fetch_url_content(url: str) -> tuple[str, str]:
    headers = {
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/avif,image/webp,image/apng,*/*;q=0.8",
        "Accept-Language": "tr-TR,tr;q=0.9,en-US;q=0.8,en;q=0.7",
        "Referer": "https://www.google.com/",
    }
    
    async with httpx.AsyncClient(timeout=TIMEOUT, verify=True, headers=headers, follow_redirects=True) as client:
        try:
            resp = await client.get(url)
            resp.raise_for_status()
            return resp.text, str(resp.url)
        except httpx.HTTPStatusError as e:
            raise HTTPException(status_code=e.response.status_code, detail=f"Site erişimi engelledi (403/404): {str(e)}")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Bağlantı hatası: {str(e)}")

def extract_main_text(html_content: str) -> str:
    extracted = trafilatura.extract(html_content, include_comments=False, include_tables=False)
    return extracted[:15000] if extracted else ""

# --- Analyze Endpoint ---
@app.post("/analyze", response_model=AnalyzeResponse)
async def analyze_news(request: AnalyzeRequest):
    clean_text = ""
    final_url = request.url

    if request.raw_text and len(request.raw_text.strip()) > 10:
        clean_text = request.raw_text[:15000]
    else:
        html_content, final_url = await fetch_url_content(request.url)
        clean_text = await run_in_threadpool(extract_main_text, html_content)

        if not clean_text or len(clean_text) < 50:
            from bs4 import BeautifulSoup
            soup = BeautifulSoup(html_content, "html.parser")
            clean_text = soup.get_text()[:5000]
    
    if len(clean_text) < 50:
        raise HTTPException(status_code=400, detail="Anlamlı metin bulunamadı / No meaningful text found.")

    # DİL TERCİHİNİ BELİRLE
    target_language = "Turkish" if request.language == "tr" else "English"
    target_audience = "Türk expat'lar" if request.language == "tr" else "expats"

    # KATEGORİLERİ DİNAMİK AL
    categories_str = ""
    if request.available_categories and len(request.available_categories) > 0:
        cat_list = ", ".join([f"'{c}'" for c in request.available_categories])
        categories_str = f"[{cat_list}]"
    else:
        categories_str = "['Genel', 'Ekonomi', 'Politika', 'Vize', 'Sosyal Yaşam']"

    # YENİ DİNAMİK PROMPT (Çift dilli sisteme uyumlu)
    system_prompt = f"""
    You are a senior news analyst creating content for {target_audience} living in Denmark.
    YOUR TASK: Analyze the news, give it a score from 1-10 based on how much it affects expats, and summarize it.
    
    CRITICAL INSTRUCTION: You MUST write your ENTIRE response (headline, source, summary, expat_significance, category) strictly in {target_language}.
    
    RULES: 
    1. Add a highly relevant emoji at the very beginning of the headline.
    2. The headline and summary must not repeat each other; provide distinct value.
    3. You must select the most appropriate category ONLY from this exact list: {categories_str}
    """

    news_schema = {
        "name": "news_analysis",
        "strict": True,
        "schema": {
            "type": "object",
            "properties": {
                "headline": {"type": "string"},
                "source": {"type": "string"},
                "summary": {"type": "string"},
                "expat_significance": {"type": "string"},
                "score": {"type": "integer"},
                "category": {"type": "string"}
            },
            "required": ["headline", "source", "summary", "expat_significance", "score", "category"],
            "additionalProperties": False,
        },
    }

    try:
        api_key_clean = request.api_key.strip()
        
        async with AsyncOpenAI(api_key=api_key_clean) as client:
            completion = await client.chat.completions.create(
                model=request.model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": f"URL: {final_url}\n\nTEXT:\n{clean_text}"},
                ],
                response_format={"type": "json_schema", "json_schema": news_schema},
            )
            data = json.loads(completion.choices[0].message.content)
            data["final_url"] = final_url
            return data
            
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"OpenAI Error: {str(e)}")

# --- Export DOCX Endpoint ---
@app.post("/export-docx")
async def export_docx(items: List[ExportItem]):
    if not items:
        raise HTTPException(status_code=400, detail="Boş liste / Empty list.")

    doc = Document()
    sorted_items = sorted(items, key=lambda x: (x.analysis.category or "Diğer", -x.analysis.score))

    for category, group_iter in groupby(sorted_items, key=lambda x: x.analysis.category or "Diğer"):
        h = doc.add_heading(category, level=1)
        run = h.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)
        
        for item in list(group_iter):
            a = item.analysis
            h2 = doc.add_heading(a.headline, level=3)
            h2.paragraph_format.space_before = Pt(12)
            
            p = doc.add_paragraph()
            p.add_run(ensure_sentence_ends(a.summary) + " ")
            
            # Word çıktısını iki dile de uyumlu olması adına "Kaynak / Source:" yaptık
            p.add_run("(Kaynak / Source: ")
            link = (a.final_url or item.input_url or "").strip()
            if link:
                add_hyperlink(p, (a.source or "Link"), link)
            else:
                p.add_run(a.source or "Link")
            p.add_run(")")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return StreamingResponse(buf, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document", headers={"Content-Disposition": 'attachment; filename="expat_intel_export.docx"'})

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=int(os.environ.get("PORT", 8000)))
